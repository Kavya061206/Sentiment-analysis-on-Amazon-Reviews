# Imports
import streamlit as st
import pandas as pd
import plotly.express as px
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from collections import Counter
import emoji
from wordcloud import WordCloud
import base64
from io import BytesIO
import io
from docx import Document
from docx.shared import Inches
import os

# Load and process data
df = pd.read_csv("E:/SentimentAnalysis/Amazon reviews/amazon_reviews.csv")
nltk.download('vader_lexicon')
sia = SentimentIntensityAnalyzer()
df['vader_score'] = df['reviewText'].apply(lambda x: sia.polarity_scores(str(x))['compound'])

def classify(score):
    if score >= 0.05:
        return 'positive'
    elif score <= -0.05:
        return 'negative'
    else:
        return 'neutral'

df['vader_sentiment'] = df['vader_score'].apply(classify)
df['vader_sentiment_clean'] = df['vader_sentiment'].astype(str).str.lower().str.strip()
df['reviewTime'] = pd.to_datetime(df['reviewTime'])

# Emoji sentiment
def extract_emojis(text):
    text = str(text)
    return ''.join(c for c in text if c in emoji.EMOJI_DATA)

emoji_sentiment_map = {
    '😊': 'positive', '😍': 'positive', '😢': 'negative',
    '😡': 'negative', '😐': 'neutral', '👍': 'positive', '👎': 'negative'
}

def classify_emoji_sentiment(emojis):
    sentiments = [emoji_sentiment_map.get(e, 'neutral') for e in emojis]
    return Counter(sentiments).most_common(1)[0][0] if sentiments else 'neutral'

df['emojis'] = df['reviewText'].apply(extract_emojis)
df['emoji_sentiment'] = df['emojis'].apply(classify_emoji_sentiment)

def hybrid_sentiment(row):
    if row['emoji_sentiment'] != row['vader_sentiment']:
        return f"{row['vader_sentiment']} (emoji: {row['emoji_sentiment']})"
    return row['vader_sentiment']

df['hybrid_sentiment'] = df.apply(hybrid_sentiment, axis=1)

# Timeline data
timeline = df.groupby(df['reviewTime'].dt.to_period('M'))['vader_sentiment'].value_counts().unstack().fillna(0)
timeline.index = timeline.index.astype(str)



# Streamlit UI
st.set_page_config(page_title="Sentiment Dashboard", layout="wide")
st.title("🩺 E-Consultation Sentiment Dashboard")

tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
    "📊 Sentiment Breakdown", "☁️ Word Cloud", "🔍 Keyword Search",
    "💬 Filtered Comments", "📈 Sentiment Timeline", "😊 Emoji Sentiment",
    "🧾 Summary", "📤 Report Generator"
])

with tab1:
    st.subheader("Sentiment Breakdown")
    sentiment_counts = df['vader_sentiment'].value_counts().reset_index()
    sentiment_counts.columns = ['Sentiment', 'Count']
    fig = px.pie(sentiment_counts, names='Sentiment', values='Count', title='Sentiment Distribution')
    st.plotly_chart(fig)

with tab2:
    st.subheader("Word Cloud")
    text = ' '.join(df['reviewText'].dropna().astype(str))
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    buffer = BytesIO()
    wordcloud.to_image().save(buffer, format="PNG")
    img_str = base64.b64encode(buffer.getvalue()).decode()
    st.markdown(f"<div style='text-align: center'><img src='data:image/png;base64,{img_str}'/></div>", unsafe_allow_html=True)

with tab3:
    st.subheader("Keyword Search")
    search_term = st.text_input("Search comments for keyword")
    if search_term:
        results = df[df['reviewText'].str.contains(search_term, case=False, na=False)]
        st.write(results[['reviewerName', 'reviewText']].head(10))

with tab4:
    st.subheader("Filtered Comments by Sentiment")
    selected_sentiment = st.selectbox("Choose sentiment", ["positive", "negative", "neutral"])
    filtered = df[df['vader_sentiment_clean'] == selected_sentiment]
    st.write(filtered[['reviewerName', 'reviewText']].head(10))

with tab5:
    st.subheader("Sentiment Timeline")
    st.line_chart(timeline)
    st.write(timeline)

with tab6:
    st.subheader("Emoji-Based Sentiment")
    st.bar_chart(df['emoji_sentiment'].value_counts())
    st.write(df[['reviewText', 'emojis', 'emoji_sentiment']].head(10))

with tab7:
    st.subheader("Sentiment Summary")
    sentiment_counts = df['vader_sentiment'].value_counts().reset_index()
    sentiment_counts.columns = ['Sentiment', 'Count']
    total = sentiment_counts['Count'].sum()
    sentiment_counts['Percentage'] = (sentiment_counts['Count'] / total * 100).round(2)
    dominant_sentiment = sentiment_counts.loc[sentiment_counts['Count'].idxmax(), 'Sentiment']
    dominant_percent = sentiment_counts.loc[sentiment_counts['Count'].idxmax(), 'Percentage']
    st.markdown(f"**Overall, the majority of public comments express a _{dominant_sentiment}_ sentiment, accounting for {dominant_percent:.2f}% of the total feedback.**")
    st.table(sentiment_counts)
with tab8:
    st.subheader("Generate Custom Word Report")

    sections = st.multiselect(
        "Select sections to include in the report:",
        ["Sentiment Breakdown", "Word Cloud", "Sentiment Timeline", "Filtered Comments", "Emoji Sentiment", "Summary"],
        default=["Summary"]
    )

    report_title = st.text_input("Enter report title:", "E-Consultation Sentiment Report")

    doc = Document()
    doc.add_heading(report_title, 0)

    # Summary section
    if "Summary" in sections:
        doc.add_heading("Summary", level=1)
        dominant = df['vader_sentiment'].value_counts().idxmax()
        percent = df['vader_sentiment'].value_counts(normalize=True).max() * 100
        doc.add_paragraph(f"Most comments are {dominant} ({percent:.2f}%).")

        sentiment_counts = df['vader_sentiment'].value_counts().reset_index()
        sentiment_counts.columns = ['Sentiment', 'Count']
        total = sentiment_counts['Count'].sum()
        sentiment_counts['Percentage'] = (sentiment_counts['Count'] / total * 100).round(2)

        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sentiment'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'

        for _, row in sentiment_counts.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Sentiment'])
            row_cells[1].text = str(row['Count'])
            row_cells[2].text = f"{row['Percentage']}%"

    # Sentiment Breakdown chart
    if "Sentiment Breakdown" in sections:
      doc.add_heading("Sentiment Breakdown", level=1)

    # Prepare data
    sentiment_counts = df['vader_sentiment'].value_counts().reset_index()
    sentiment_counts.columns = ['Sentiment', 'Count']

    # Create pie chart using Plotly Express
    import plotly.express as px
    pie_fig = px.pie(
        sentiment_counts,
        names='Sentiment',
        values='Count',
        title='Sentiment Distribution',
        color_discrete_sequence=px.colors.qualitative.Set2  # Matches dashboard colors
    )

    # Set layout to match dashboard
    pie_fig.update_layout(
        paper_bgcolor='white',
        plot_bgcolor='white',
        font=dict(color='black'),
        showlegend=True,
        legend=dict(
            orientation="v",
            x=1.05,
            y=0.5
        )
    )

    # Save image
    pie_fig.write_image("sentiment_pie.png", width=800, height=600, scale=2)
    doc.add_picture("sentiment_pie.png", width=Inches(5.5))



    # Word Cloud
    if "Word Cloud" in sections:
        doc.add_heading("Word Cloud", level=1)
        text = ' '.join(df['reviewText'].dropna().astype(str))
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
        wordcloud_image = wordcloud.to_image()
        wordcloud_image.save("wordcloud.png")
        doc.add_picture("wordcloud.png", width=Inches(5.5))

    # Sentiment Timeline
    if "Sentiment Timeline" in sections:
        doc.add_heading("Sentiment Timeline", level=1)
        timeline = df.groupby(df['reviewTime'].dt.to_period('M'))['vader_sentiment'].value_counts().unstack().fillna(0)
        timeline.index = timeline.index.astype(str)
        fig = px.line(timeline)
        fig.update_layout(paper_bgcolor='white', plot_bgcolor='white')
        fig.write_image("timeline_chart.png")
        doc.add_picture("timeline_chart.png", width=Inches(5.5))

    # Emoji Sentiment
    if "Emoji Sentiment" in sections:
        doc.add_heading("Emoji Sentiment", level=1)
        emoji_summary = df['emoji_sentiment'].value_counts().to_string()
        doc.add_paragraph(emoji_summary)

    # Filtered Comments
    if "Filtered Comments" in sections:
        doc.add_heading("Sample Comments", level=1)
        sample_comments = df[['reviewText', 'vader_sentiment']].head(5)
        for _, row in sample_comments.iterrows():
            doc.add_paragraph(f"{row['vader_sentiment'].capitalize()}: {row['reviewText']}", style='List Bullet')

    # Save and download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Word Report",
        data=buffer,
        file_name="sentiment_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Clean up images
    for img in ["sentiment_pie.png", "wordcloud.png", "timeline_chart.png"]:
        if os.path.exists(img):
            os.remove(img)
